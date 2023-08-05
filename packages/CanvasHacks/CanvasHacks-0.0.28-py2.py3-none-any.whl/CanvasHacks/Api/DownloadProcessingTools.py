"""
Created by adam on 1/29/19
"""
import json
import io
import PyPDF2
import docx
import requests
from requests.exceptions import HTTPError

from CanvasHacks.Api.RequestTools import make_request_header
from CanvasHacks.Api.UrlTools import make_url

__author__ = 'adam'

PICTURE_PLACEHOLDER = 'picture-uploaded picture-uploaded picture-uploaded picture-uploaded picture-uploaded '


def download_journals( url, destination ):
    """Handles making the request to download journals"""
    # download file
    print( url )
    file_response = requests.get( url, headers=make_request_header() )

    # write to folder
    print( "Writing to %s" % destination )
    open( destination, 'wb' ).write( file_response.content )


def download_submitted_file( url, filepath ):
    response = requests.get( url, headers=make_request_header() )
    # save it to file
    open( filepath, 'wb' ).write( response.content )


def get_submissions( course_id, assignment_id, assign_type='assignments', per_page=42 ):
    """Makes request to the server for all submissions for the given unit
    Example
        course_id = SECTION_930
        assignment_id = 288480
        response2 = get_submissions(course_id, assignment_id)
    """
    responses = [ ]
    url = make_url( course_id, assign_type )
    url = "%s/%s/submissions?per_page=%s" % (url, assignment_id, per_page)
    try:
        while True:
            print( url )
            response = requests.get( url, headers=make_request_header() )
            # If the response was successful, no Exception will be raised
            response.raise_for_status()
            # Continuing on since was successful
            responses += response.json()
            url = response.links[ 'next' ][ 'url' ]

    except KeyError:
        return responses

    except HTTPError as http_err:
        print(f'HTTP error occurred: {http_err}')  # Python 3.6
        return responses

    except Exception as err:
        print(f'Other error occurred: {err}')  # Python 3.6
        return responses





def process_response( response_json, journal_folder=None ):
    """Takes the response and pulls out submissions which used the text box, then downloads
    submitted files and processes out their text content
    """
    submissions = [ ]

    for j in response_json:
        result = { 'submission_id': j[ 'id' ], 'student_id': j[ 'user_id' ] }
        if j[ 'body' ] is not None:
            # The student used the online text box to submit
            result[ 'body' ] = j[ 'body' ]

        else:
            # The student submitted the journal as a separate document
            if 'attachments' in j.keys() and len( j[ 'attachments' ] ) > 0:
                url = j[ 'attachments' ][ 0 ][ 'url' ]
                filename = make_journal_filename( j )

                # download and save submitted file
                fpath = "%s/%s" % (journal_folder, filename)
                download_submitted_file( url, fpath )

                # open the file and extract text
                result[ 'body' ] = getBody( fpath )

            else:
                # NB., if a student never submitted (workflow_state = 'unsubmitted'),
                # then they will have an entry which lacks a body key
                result[ 'body' ] = None
        submissions.append( result )

    return submissions


def save_submission_json( submissions, folder, json_name='all-submissions' ):
    # save submissions
    with open( "%s/%s.json" % (folder, json_name), 'w' ) as fpp:
        json.dump( submissions, fpp )


def get_journal_filename( response ):
    """Extracts the name of the file submitted from the response"""
    return response[ 'attachments' ][ 0 ][ 'filename' ]


def make_journal_filename( response ):
    """Creates the standardized filename for saving"""
    return "%s-%s" % (response[ 'user_id' ], get_journal_filename( response ))


# -------------- Document processing handlers

def docx_handler( filename ):
    if type(filename) is bytes:
        doc = docx.Document(io.BytesIO(filename))
    else:
        doc = docx.Document( filename )
    fullText = [ ]
    for para in doc.paragraphs:
        fullText.append( para.text )

    return '\n'.join( fullText )


def pdf_handler( filepath ):
    """Extracts the text from a pdf file"""
    if type(filepath) is bytes:
        # if using online version and not saving
        # downloads to file
        pdfReader = PyPDF2.PdfFileReader( io.BytesIO(filepath))
    else:
        # If loading downloaded work from file
        pdfFileObj = open( filepath, 'rb' )
        pdfReader = PyPDF2.PdfFileReader( pdfFileObj )
    pageObj = pdfReader.getPage( 0 )
    return pageObj.extractText()


def picture_handler( v=None ):
    return PICTURE_PLACEHOLDER


def unknown_handler( v=None ):
    return 'u'
    # print( 'unknown', v)
    return None


handlers = {
    '.docx': docx_handler,
    '.doc': unknown_handler,
    '.pdf': pdf_handler,
    '.jpg': picture_handler,
    '.png': picture_handler,
    '.xml': picture_handler,
    '.pages': picture_handler,
    '.odt': picture_handler
}


def chooseHandler( url ):
    """Checks the end of the file download url to determine
    whether it matches any of the file types we have handlers
    registered for. Returns the relevant handler"""
    url = url.strip()
    lengths = [len(a) for a in handlers.keys()]
    for i in range( min(lengths), max(lengths) + 1):
        if url[ -i: ] in handlers.keys():
            return handlers.get( url[ -i: ])

    return unknown_handler


def getBody( filepath ):
    """Extracts the text from a saved file"""
    PICTURE_PLACEHOLDER = 'picture-uploaded picture-uploaded picture-uploaded picture-uploaded picture-uploaded '
    filepath = filepath.strip()

    if filepath[ -5: ] == '.docx':
        return docx_handler( filepath )

    if filepath[ -4: ] == '.pdf':
        return pdf_handler( filepath )

    elif filepath[ -4: ] == '.jpg':
        return PICTURE_PLACEHOLDER + filepath[ -4: ]

    elif filepath[ -4: ] == '.png':
        return PICTURE_PLACEHOLDER + filepath[ -4: ]

    elif filepath[ -4: ] == '.xml':
        return PICTURE_PLACEHOLDER + filepath[ -4: ]

    elif filepath[ -6: ] == '.pages':
        return PICTURE_PLACEHOLDER + filepath[ -6: ]

    elif filepath[ -4: ] == '.odt':
        return PICTURE_PLACEHOLDER + filepath[ -4: ]

    else:
        return PICTURE_PLACEHOLDER + filepath


def process_response_without_saving_files( response_json ):
    """Takes the response and pulls out submissions which used the text box, then downloads
    submitted files and processes out their text content
    """
    submissions = [ ]

    for j in response_json:
        result = { 'submission_id': j[ 'id' ], 'student_id': j[ 'user_id' ] }
        if j[ 'body' ] is not None:
            # The student used the online text box to submit
            result[ 'body' ] = j[ 'body' ]

        else:
            # The student submitted the journal as a separate document
            if 'attachments' in j.keys() and len( j[ 'attachments' ] ) > 0:
                url = j[ 'attachments' ][ 0 ][ 'url' ]
                # download the submitted file
                # print( "Downloading: ", url )
                response = requests.get( url, headers=make_request_header() )
                content = response.content

                # determine the appropriate handler to use
                filename = response.headers['content-disposition'].split('filename=')[1]
                filename = filename.replace('"', '')
                handler = chooseHandler(filename)

                # open the file and extract text
                result[ 'body' ] = handler( content )

            else:
                # NB., if a student never submitted (workflow_state = 'unsubmitted'),
                # then they will have an entry which lacks a body key
                result[ 'body' ] = None
        submissions.append( result )

    return submissions


def extract_body(submission):
    """Given a canvas api object returns the body text for the submission"""
    if submission.body is not None:
        # The student used the online text box to submit
        return submission.body

    else:
        # The student submitted the journal as a separate document
        if 'attachments' in submission.attributes.keys() and len( submission.attributes[ 'attachments' ] ) > 0:
            url = submission.attachments[ 0 ][ 'url' ]
            # download the submitted file
            # print( "Downloading: ", url )
            response = requests.get( url, headers=make_request_header() )
            content = response.content

            # determine the appropriate handler to use
            filename = response.headers['content-disposition'].split('filename=')[1]
            filename = filename.replace('"', '')
            handler = chooseHandler(filename)

            # open the file and extract text
            return handler( content )

        else:
            # NB., if a student never submitted (workflow_state = 'unsubmitted'),
            # then they will have an entry which lacks a body key
            return None



if __name__ == '__main__':
    pass
